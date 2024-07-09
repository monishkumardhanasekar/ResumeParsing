from docx import Document
import fitz

def extract_text_from_pdf(file):
    text = ""
    pdf_document = fitz.open(stream=file.read(), filetype="pdf")  # Open PDF from file-like object
    for page in pdf_document:
        text += page.get_text()  # Extract text from each page
    return text

def extract_text_from_docx(file):
    """
    Extracts text from a DOCX file, including text from tables.
    
    :param docx_path: Path to the DOCX file.
    :return: Extracted text as a single string.
    """
    text = ""
    try:
        # Open the DOCX file
        doc = Document(file)
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
    except Exception as e:
        print(f"An error occurred while extracting text from DOCX: {e}")
    return text


def extract_text_from_txt(file):
    text = file.read().decode('utf-8')  # Read the text file and decode it as UTF-8
    return text